#!/usr/bin/env python3
"""
Generate PRIVACY_POLICY.docx and TERMS_OF_SERVICE.docx from markdown
with Times New Roman 11pt font.
Requires: pip install python-docx
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def set_document_font(doc, font_name='Times New Roman', font_size_pt=11):
    """Set default font for body and headings."""
    size = Pt(font_size_pt)
    for style in doc.styles:
        if hasattr(style, 'font') and style.font is not None:
            try:
                style.font.name = font_name
                style.font.size = size
            except Exception:
                pass


def md_to_docx(md_path, docx_path):
    """Convert markdown file to docx with Times New Roman 11pt."""
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    text = md_path.read_text(encoding='utf-8')
    doc = Document()
    set_document_font(doc, 'Times New Roman', 11)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == '---':
            i += 1
            continue

        # Title: first ## (e.g. "## Privacy Policy")
        if stripped.startswith('## ') and not re.match(r'^##\s+\d+\.', stripped) and not re.match(r'^##\s+\d+\.\d+', stripped):
            title = stripped[3:].strip()
            h = doc.add_heading(title, level=0)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            i += 1
            continue

        # Section heading: ## 1. Section or ## 9.1. Section
        if re.match(r'^##\s+[\d.]+', stripped):
            title = re.sub(r'^##\s+', '', stripped).strip()
            h = doc.add_heading(title, level=1)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            i += 1
            continue

        # List item: - **Bold** or - text
        if stripped.startswith('- '):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            parts = re.split(r'(\*\*[^*]+\*\*)', item_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
            i += 1
            while i < len(lines) and len(lines[i]) > 0 and lines[i].startswith('  ') and lines[i].strip().startswith('- '):
                sub = lines[i].strip()[2:].strip()
                p = doc.add_paragraph(style='List Bullet 2')
                p.paragraph_format.space_after = Pt(2)
                parts = re.split(r'(\*\*[^*]+\*\*)', sub)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        r = p.add_run(part[2:-2])
                        r.bold = True
                    else:
                        r = p.add_run(part)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
                i += 1
            continue

        # Italic-only line (e.g. _Last updated_)
        if stripped.startswith('_') and stripped.endswith('_'):
            p = doc.add_paragraph()
            r = p.add_run(stripped[1:-1])
            r.italic = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Plain paragraph
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2] + ' ')
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
            i += 1
            continue

        i += 1

    doc.save(docx_path)
    print(f"Created {docx_path}")


def main():
    base = Path(__file__).resolve().parent.parent
    md_to_docx(base / 'PRIVACY_POLICY.md', base / 'PRIVACY_POLICY.docx')
    md_to_docx(base / 'TERMS_OF_SERVICE.md', base / 'TERMS_OF_SERVICE.docx')


if __name__ == '__main__':
    main()
